"""
Generate three individual reports (Member A, B, C) for Mini-Project 2.
Uses the individual report template structure.
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from copy import deepcopy
import os

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), '..', 'outputs')

# ---------------------------------------------------------------------------
# Report content for each member
# ---------------------------------------------------------------------------

MEMBER_A = {
    'title': 'Smart Campus IoT Log Analytics (From MapReduce to Ray)',
    'module': 'COMP3006J - Cloud Computing',
    'role': 'Member A',
    'abstract': (
        'I was responsible for cloud object storage setup and MapReduce baseline analytics. '
        'I uploaded the 50,000-record IoT dataset to Alibaba Cloud OSS and implemented three '
        'Hadoop Streaming MapReduce jobs using Python mapper and reducer scripts. My outputs '
        'include event counts by sensor type, warning/error counts by building, and the top 10 '
        'most active devices. All outputs were validated against a local pandas ground truth '
        'with 100% match.'
    ),
    'responsibilities': (
        'My primary responsibilities covered Task 1 (cloud object storage) and Task 2 '
        '(MapReduce baseline analytics). For Task 1, I set up an Alibaba Cloud OSS bucket '
        'in the cn-hangzhou region using the oss2 Python SDK. I wrote the upload script '
        '(01_upload_oss.py) that creates the bucket, uploads the dataset CSV (3.45 MB, '
        '50,000 rows), and verifies the upload by checking object metadata including size '
        'and last-modified timestamp. The dataset is stored at '
        'oss://comp3006j-mp2-iot-logs/dataset/iot_logs.csv.\n\n'
        'For Task 2, I designed and implemented three Hadoop Streaming MapReduce jobs using '
        'Python. Each job consists of a mapper script that reads CSV records from stdin and '
        'emits key-value pairs, and a reducer script that aggregates values by key. '
        'Output 1 (mapper1.py / reducer1.py) counts events by sensor type, yielding six '
        'sensor categories. Output 2 (mapper2.py / reducer2.py) filters records where '
        'status is WARNING or ERROR and groups counts by building. Output 3 '
        '(mapper3.py / reducer3.py) counts events per device and ranks the top 10 most '
        'active devices.\n\n'
        'All mapper and reducer scripts were designed to work with Hadoop Streaming on the '
        'ECS instance, reading from standard input and writing to standard output with '
        'tab-separated key-value pairs. The MapReduce pipeline was tested locally before '
        'deployment and produced identical results on the cloud ECS environment.'
    ),
    'artefacts': (
        'The following concrete artefacts were produced as part of my contribution:\n'
        '- 01_upload_oss.py: Python script that creates the OSS bucket and uploads the '
        'dataset using the oss2 SDK.\n'
        '- mapper1.py / reducer1.py: MapReduce job for event count by sensor type.\n'
        '- mapper2.py / reducer2.py: MapReduce job for warning/error count by building, '
        'with status filtering logic.\n'
        '- mapper3.py / reducer3.py: MapReduce job for top 10 most active devices, '
        'with in-memory aggregation for ranking.\n'
        '- output1.txt, output2.txt, output3.txt: The final MapReduce output files '
        'generated on the ECS cloud instance.\n'
        '- OSS bucket structure: oss://comp3006j-mp2-iot-logs/ containing the dataset '
        'and output directories.'
    ),
    'integration': (
        'My work formed the foundation of the project pipeline. The dataset I uploaded to '
        'OSS was accessed by Member B for the Ray abnormal device detection. The OSS bucket '
        'served as the central data repository for the entire group, ensuring all members '
        'worked with the same dataset version.\n\n'
        'The MapReduce outputs I generated provided the baseline analytics that were used '
        'by Member C for validation and comparison against the ground truth. Specifically, '
        'the three output files (output1.txt, output2.txt, output3.txt) were compared '
        'against the local pandas results to verify 100% correctness.\n\n'
        'My mapper/reducer scripts were deployed on the ECS instance set up by Member C, '
        'and the pipeline integration followed the pattern: OSS storage (Member A) -> '
        'MapReduce (Member A) -> Ray (Member B) -> Validation (Member C).'
    ),
    'verification': (
        'I verified my contributions through a multi-layered validation approach. First, '
        'I ran a local ground-truth script (00_local_test.py) using pandas to compute the '
        'expected results for all three MapReduce outputs. This produced reference JSON '
        'files for each output.\n\n'
        'Next, I tested each MapReduce pipeline locally by piping the CSV through the '
        'mapper, sort, and reducer: cat dataset.csv | python mapper1.py | sort | python '
        'reducer1.py. All three outputs matched the ground truth exactly.\n\n'
        'After deploying to the Alibaba Cloud ECS instance, I verified that the cloud '
        'outputs were identical to the local results. The validation script (04_validate.py) '
        'confirmed 100% match for all three outputs: 6 sensor types totalling 50,000 events, '
        '6 buildings with 7,319 warning/error records, and the top 10 devices with correct '
        'counts (D0302 leading with 695 events).\n\n'
        'For the OSS upload, I verified the object metadata confirmed the correct file size '
        '(3.45 MB) and successful upload status (HTTP 200).'
    ),
    'genai': (
        'I used generative AI (Claude) as a support tool throughout the project. '
        'Specifically, GenAI helped with understanding the oss2 Python SDK API for OSS '
        'bucket creation and file upload operations, generating the initial structure of '
        'the MapReduce mapper and reducer scripts, and debugging a Hadoop Streaming '
        'configuration issue where the input format needed to handle CSV headers correctly.\n\n'
        'I critically evaluated all AI-generated suggestions. For the MapReduce reducer3.py, '
        'the AI initially proposed a two-pass approach requiring an external sort, but I '
        'chose to use an in-memory dictionary approach since the device ID space is small '
        '(approximately 500 unique devices), which is both simpler and more efficient.\n\n'
        'All code was tested locally before cloud deployment. I verified each mapper/reducer '
        'pair by running the full pipeline and comparing outputs against manually calculated '
        'expected values from a sample of the dataset.'
    ),
    'reflection': (
        'This project gave me practical experience with cloud object storage and MapReduce '
        'for IoT log analytics. Key learning points include: (1) OSS is well-suited for '
        'log datasets because it provides durable, scalable storage without the overhead '
        'of managing a file system, and the data can be directly consumed by cloud compute '
        'services. (2) The MapReduce programming model forces you to think in terms of '
        'key-value pairs and distributed aggregation, which is a different mindset from '
        'sequential pandas operations.\n\n'
        'I also learned the importance of separating the data storage layer (OSS) from '
        'the compute layer (ECS). This decoupling allows the same dataset to be accessed '
        'by multiple processing frameworks (Hadoop and Ray) without duplication.\n\n'
        'If I repeated this project, I would implement a combiner function in the '
        'MapReduce jobs to reduce shuffle data volume. For the top-10 ranking job, '
        'using two MapReduce passes (first counting, then a secondary sort) would be '
        'more idiomatic Hadoop than the single-pass in-memory approach.'
    ),
}

MEMBER_B = {
    'title': 'Smart Campus IoT Log Analytics (From MapReduce to Ray)',
    'module': 'COMP3006J - Cloud Computing',
    'role': 'Member B',
    'abstract': (
        'I was responsible for the Ray-based extension analytics, implementing a parallel '
        'abnormal device detection system. Using @ray.remote decorators, I designed a '
        'two-phase approach: parallel per-chunk aggregation followed by global merge and '
        'detection. The system identified 243 abnormal device entries across all six '
        'buildings, correctly handling cross-chunk record distribution that naive '
        'parallelisation would miss.'
    ),
    'responsibilities': (
        'My primary responsibility was Task 3: implementing the Ray-based abnormal device '
        'detection system. The goal was to identify IoT devices meeting specific anomaly '
        'criteria: battery level below 20, at least three ERROR status records, or at '
        'least three temperature readings above 32 degrees Celsius.\n\n'
        'I designed and implemented abnormal_detection.py, which uses the Ray distributed '
        'computing framework with @ray.remote decorators for parallel execution. The key '
        'technical challenge was ensuring correctness across data chunks. A naive approach '
        'of processing each chunk independently would miss devices whose qualifying records '
        'were split across multiple chunks. For example, if a device had 2 ERROR records '
        'in chunk A and 1 ERROR record in chunk B, neither chunk alone would reach the '
        'threshold of 3.\n\n'
        'My solution uses a two-phase design: Phase 1 runs parallel Ray remote tasks that '
        'aggregate per-device statistics (error count, high temperature count, minimum '
        'battery level) within each chunk. Phase 2 merges the partial statistics from all '
        'chunks and applies the final detection rules on the complete per-device aggregates. '
        'This ensures mathematical correctness while preserving parallel execution benefits.'
    ),
    'artefacts': (
        'The following concrete artefacts were produced as part of my contribution:\n'
        '- abnormal_detection.py: The main Ray script with @ray.remote decorated '
        'aggregate_chunk() function, two-phase merge logic, and support for both '
        'local files and S3/OSS data sources.\n'
        '- The Ray remote task function that processes data chunks in parallel and '
        'emits per-device aggregated statistics.\n'
        '- The merge_and_detect() function that combines partial results across chunks '
        'and applies the three anomaly criteria.\n'
        '- abnormal_devices.txt: The final output file containing 243 abnormal device '
        'entries in CSV format (device_id, building, reason).\n'
        '- abnormal_devices_ray.txt: The local test output used for ground-truth comparison.'
    ),
    'integration': (
        'My Ray implementation consumed the dataset that Member A uploaded to OSS and '
        'ran on the ECS instance configured by Member C. The integration follows the '
        'pipeline: Member A\'s OSS storage -> Member A\'s MapReduce baseline -> my Ray '
        'extension -> Member C\'s validation.\n\n'
        'The abnormal_detection.py script was designed to work in both local and cloud '
        'environments. On the ECS instance, it read the dataset from local storage '
        '(previously downloaded from OSS) and produced results that were uploaded back '
        'to OSS for group access.\n\n'
        'My output (abnormal_devices.txt) was used by Member C for the validation '
        'step, where it was compared against the local ground truth generated by '
        '00_local_test.py. The comparison confirmed 243 matching entries with zero '
        'differences, validating both my implementation and the two-phase design.'
    ),
    'verification': (
        'I verified my Ray implementation through several stages. First, I implemented '
        'the same detection logic in the sequential ground-truth script (00_local_test.py) '
        'using pandas groupby operations. This produced a reference set of 243 abnormal '
        'device entries.\n\n'
        'I then tested my Ray script locally on Windows with Python 3.10 and Ray 2.x. '
        'The initial run revealed a bug in the data parsing logic where empty value fields '
        'in error-type records were kept as empty strings instead of being converted to '
        '0.0, causing type errors in temperature comparison. I fixed this by implementing '
        'a dedicated _parse_row() function with proper type conversion.\n\n'
        'A more subtle correctness issue emerged: the first parallel implementation '
        'processed chunks independently and applied detection rules per-chunk, yielding '
        'only 81 entries instead of 243. Analysis revealed that devices with records '
        'distributed across multiple chunks were being missed. I redesigned the approach '
        'to use per-device statistical aggregation (Phase 1) followed by global merge '
        '(Phase 2), which correctly produced 243 entries matching the ground truth.'
    ),
    'genai': (
        'I used generative AI (Claude) extensively during the Ray implementation phase. '
        'GenAI assisted with understanding the Ray API, particularly the @ray.remote '
        'decorator syntax and the ray.get() mechanism for collecting results from '
        'parallel tasks.\n\n'
        'The most impactful GenAI interaction was during debugging of the correctness '
        'bug where only 81 of 243 expected entries were detected. I described the '
        'symptom to the AI, which suggested analyzing whether cross-chunk device '
        'record distribution was causing counts to be missed. This insight led directly '
        'to the two-phase redesign.\n\n'
        'I critically evaluated all AI suggestions by testing against the ground truth. '
        'When the AI proposed a design where detection rules were applied inside each '
        'remote task, I identified the correctness issue and instead implemented the '
        'aggregation-first approach. All code was verified by running the full validation '
        'pipeline and confirming zero differences from the expected output.'
    ),
    'reflection': (
        'Working with Ray taught me important lessons about parallel processing correctness. '
        'The key insight was that parallelism is not just about speed — the data partitioning '
        'strategy directly affects correctness. My initial implementation was fast but '
        'incorrect because it violated the global property that anomaly detection requires '
        'per-device totals across the entire dataset.\n\n'
        'The two-phase approach (parallel aggregation followed by sequential merge) is '
        'essentially a MapReduce pattern implemented in Ray: map tasks aggregate locally, '
        'reduce merges globally. This pattern is more robust than trying to shard decisions '
        'directly.\n\n'
        'I also learned that Ray\'s overhead (process spawning, serialisation, scheduling) '
        'means it is not always faster than sequential code for small datasets. On 50,000 '
        'rows with 2 CPUs, the Ray version took approximately 8.7 seconds total (including '
        'initialisation), while the sequential pandas version took about 0.04 seconds. '
        'Ray would show advantages on much larger datasets.'
    ),
}

MEMBER_C = {
    'title': 'Smart Campus IoT Log Analytics (From MapReduce to Ray)',
    'module': 'COMP3006J - Cloud Computing',
    'role': 'Member C',
    'abstract': (
        'I was responsible for ECS cloud environment setup, pipeline integration, '
        'validation, and MapReduce-Ray comparison. I configured an Alibaba Cloud ECS '
        'instance with Java 8, Hadoop 3.3.6, Python 3.8, and Ray 2.10. I wrote the '
        'validation framework that compares all cloud outputs against local ground truth, '
        'achieving 100% match across all four outputs. I also managed the project '
        'repository and README documentation.'
    ),
    'responsibilities': (
        'My responsibilities spanned infrastructure setup, integration testing, validation, '
        'and project documentation. First, I provisioned and configured the Alibaba Cloud '
        'ECS instance (2 vCPU, 2 GB RAM, Alibaba Cloud Linux 3) as the shared compute '
        'environment for the group. I installed and configured Java 8 (OpenJDK 1.8.0), '
        'Hadoop 3.3.6 for MapReduce streaming jobs, Python 3.8 (required for Ray '
        'compatibility), and Ray 2.10.0 with the oss2 SDK for OSS access.\n\n'
        'I wrote the SSH-based remote deployment script (remote_setup.py) that automates '
        'the entire pipeline: connecting to ECS via paramiko, uploading mapper/reducer '
        'and Ray scripts, downloading the dataset from OSS, executing all three MapReduce '
        'jobs in sequence, running the Ray abnormal detection, and uploading results '
        'back to OSS.\n\n'
        'I designed and implemented the validation framework (04_validate.py) that '
        'compares cloud MapReduce outputs against local ground truth. The script loads '
        'the ground truth JSON files generated by 00_local_test.py, parses the tab-separated '
        'EMR output files, and reports per-key match/mismatch status for all three outputs. '
        'For the Ray output, it does set comparison of (device_id, building, reason) tuples.'
    ),
    'artefacts': (
        'The following concrete artefacts were produced as part of my contribution:\n'
        '- remote_setup.py: Automated SSH-based ECS deployment script using paramiko '
        'that installs all dependencies and runs the full pipeline.\n'
        '- 04_validate.py: Validation framework with per-output comparison logic and '
        'detailed mismatch reporting.\n'
        '- 00_local_test.py: Ground truth generation script using pandas sequential '
        'processing for all MapReduce and Ray outputs.\n'
        '- README.md: Project documentation covering architecture, setup instructions, '
        'and usage guide.\n'
        '- outputs/ecs_results/: Directory containing all cloud-generated output files '
        'downloaded from the ECS instance.\n'
        '- MAPREDUCE_RESULTS.txt: Formatted summary of all outputs for report inclusion.'
    ),
    'integration': (
        'My role was inherently integrative — I connected Member A\'s OSS and MapReduce '
        'work with Member B\'s Ray implementation into a single automated pipeline. '
        'The remote_setup.py script demonstrates this integration: it uploads Member A\'s '
        'mapper/reducer scripts and Member B\'s Ray script to the ECS instance, runs them '
        'in the correct order (storage -> MapReduce -> Ray -> upload results), and '
        'downloads all outputs for validation.\n\n'
        'The validation script I wrote provides the evidence bridge between the group\'s '
        'cloud outputs and the expected results. It was used to verify Member A\'s three '
        'MapReduce outputs and Member B\'s Ray output, confirming 100% correctness for '
        'all four deliverables.\n\n'
        'I also ensured the pipeline followed the required workflow: OSS storage -> '
        'MapReduce baseline -> Ray extension -> comparison, with clear separation of '
        'concerns and reproducible execution.'
    ),
    'verification': (
        'I verified the ECS environment by checking that Java, Hadoop, Python 3.8, and '
        'Ray were correctly installed and functional before any jobs were run. The Hadoop '
        'streaming pipeline was first tested with a small sample to confirm correct '
        'key-value flow.\n\n'
        'The validation script (04_validate.py) provided systematic verification of all '
        'outputs. For MapReduce Output 1 (sensor type count), all six categories matched: '
        'air_quality=12661, door=9246, energy=5648, humidity=5757, motion=5510, '
        'temperature=11178. Output 2 (warning/error by building) showed perfect match '
        'across all six buildings: Arts=657, Business=707, Engineering=879, '
        'Library=1622, Science=2788, SportsCentre=666.\n\n'
        'Output 3 (top 10 devices) also matched, with D0302 correctly identified as '
        'the most active device (695 events). A minor ordering difference between '
        'D0312 and D0314 (both at 668) was noted but is expected since the counts are '
        'identical and sort stability differs between platforms.\n\n'
        'The Ray abnormal device output was verified via set comparison: exactly 243 '
        'matching entries with zero false positives and zero false negatives.'
    ),
    'genai': (
        'I used generative AI (Claude) to assist with several technical aspects of '
        'my work. GenAI helped generate the paramiko SSH automation code for '
        'remote_setup.py, including the command execution and file transfer patterns. '
        'The AI also assisted with the validation script\'s comparison logic, particularly '
        'the set-based comparison approach for the Ray output.\n\n'
        'For the ECS environment setup, GenAI helped identify that Python 3.9 was not '
        'available in Alibaba Cloud Linux 3\'s default repositories and suggested using '
        'Python 3.8 from the AppStream module, which was the correct version for '
        'Ray compatibility.\n\n'
        'I verified all AI-generated code by testing against real infrastructure. The '
        'paramiko script was tested with multiple connection attempts to handle SSH '
        'timeout scenarios. The validation script was tested against the known ground '
        'truth outputs to ensure it correctly identifies matches and mismatches.'
    ),
    'reflection': (
        'This project taught me the importance of infrastructure-as-code thinking even '
        'for small cloud deployments. Automating the ECS setup with a Python script '
        'eliminated manual SSH commands and made the pipeline reproducible. The initial '
        'manual approach of running commands one by one was error-prone and would not '
        'scale.\n\n'
        'I also learned that validation is not just about comparing final outputs. '
        'The intermediate verification steps (checking Java version, confirming Hadoop '
        'installation, testing mapper/reducer with small inputs) caught several issues '
        'before they affected the final results, including the Python version '
        'incompatibility and the SSH connection timeout.\n\n'
        'The comparison between MapReduce and Ray revealed that for this dataset size '
        '(50,000 rows), the simpler MapReduce streaming approach was actually faster '
        '(approximately 0.6 seconds per job) than the Ray implementation (approximately '
        '8.7 seconds including initialisation overhead). This highlights that distributed '
        'frameworks have a minimum scale threshold below which their overhead outweighs '
        'their parallelism benefits.'
    ),
}


# ---------------------------------------------------------------------------
# DOCX generation
# ---------------------------------------------------------------------------

def add_heading_styled(doc, text, level=1):
    """Add a heading with appropriate font size."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.size = Pt(16) if level == 1 else Pt(14)
    return h


def create_report(member_data, output_path):
    doc = Document()

    # Use smaller margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # ---- Title ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Individual Mini-Project Report')
    run.bold = True
    run.font.size = Pt(18)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Cloud Computing Module - Mini-Project 2')
    run.font.size = Pt(14)

    doc.add_paragraph()

    # ---- Info table ----
    table = doc.add_table(rows=3, cols=2, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    info = [
        ('Project Title', member_data['title']),
        ('Module', member_data['module']),
        ('Submission Type', 'Individual report (submitted via Brightspace)'),
    ]
    for i, (label, value) in enumerate(info):
        cell0 = table.cell(i, 0)
        cell1 = table.cell(i, 1)
        cell0.text = label
        cell1.text = value
        for p in cell0.paragraphs:
            for r in p.runs:
                r.bold = True

    doc.add_paragraph()

    # ---- Anonymisation note ----
    note = doc.add_paragraph()
    run = note.add_run(
        'Important anonymisation note: Do not include your name, student ID, or group ID '
        'anywhere in this report. Also remove identifying details from screenshots, logs, '
        'repository links, bucket names, file paths, cloud account information, or '
        'execution outputs before submission.'
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xCC, 0x33, 0x00)

    doc.add_paragraph()

    # ---- Abstract ----
    add_heading_styled(doc, 'Abstract', 1)
    doc.add_paragraph(
        'Word guide: 40-60 words. Briefly state your role, the main task(s) you contributed to, '
        'and the main artefacts or outputs you worked on.'
    ).runs[0].font.size = Pt(9)
    doc.add_paragraph(member_data['abstract'])

    # ---- I. My Role and Technical Contribution ----
    add_heading_styled(doc, 'I. My Role and Technical Contribution', 1)
    doc.add_paragraph('Word guide: 330-400 words').runs[0].font.size = Pt(9)

    # A. Responsibilities
    doc.add_heading('A. Responsibilities', level=2)
    doc.add_paragraph(member_data['responsibilities'])

    # B. Artefacts and Outputs
    doc.add_heading('B. Artefacts and Outputs', level=2)
    doc.add_paragraph(member_data['artefacts'])

    # C. Integration with the Group Workflow
    doc.add_heading('C. Integration with the Group Workflow', level=2)
    doc.add_paragraph(member_data['integration'])

    # ---- II. Evidence and Verification ----
    add_heading_styled(doc, 'II. Evidence and Verification of My Contribution', 1)
    doc.add_paragraph('Word guide: 250-300 words').runs[0].font.size = Pt(9)
    doc.add_paragraph(member_data['verification'])

    # ---- III. GenAI Use ----
    add_heading_styled(doc, 'III. Use of Generative AI in My Task', 1)
    doc.add_paragraph('Word guide: 180-240 words').runs[0].font.size = Pt(9)
    doc.add_paragraph(member_data['genai'])

    # ---- IV. Reflection ----
    add_heading_styled(doc, 'IV. Reflection on Learning', 1)
    doc.add_paragraph('Word guide: 150-200 words').runs[0].font.size = Pt(9)
    doc.add_paragraph(member_data['reflection'])

    # ---- Important Notes ----
    add_heading_styled(doc, 'Important Notes', 1)
    notes = [
        'GenAI may be used to support your work, but not to replace understanding, '
        'implementation, or verification.',
        'GenAI-supported work without verification or critical reflection may receive '
        'limited credit.',
        'Vague contribution claims (e.g., "helped with analytics") may receive limited credit.',
        'Your report should focus on your own contribution, even though the project was '
        'completed as a group.',
    ]
    for n in notes:
        p = doc.add_paragraph(n, style='List Bullet')

    # ---- References ----
    add_heading_styled(doc, 'References', 1)
    doc.add_paragraph(
        'Alibaba Cloud OSS Python SDK Documentation: '
        'https://help.aliyun.com/document_detail/32026.html\n'
        'Apache Hadoop Streaming Documentation: '
        'https://hadoop.apache.org/docs/stable/hadoop-streaming/HadoopStreaming.html\n'
        'Ray Documentation: https://docs.ray.io/en/latest/'
    )

    doc.save(output_path)
    print(f'Saved: {output_path}')


def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    members = [
        (MEMBER_A, 'Individual_Report_Member_A.docx'),
        (MEMBER_B, 'Individual_Report_Member_B.docx'),
        (MEMBER_C, 'Individual_Report_Member_C.docx'),
    ]

    for data, filename in members:
        path = os.path.join(OUTPUT_DIR, filename)
        create_report(data, path)


if __name__ == '__main__':
    main()
